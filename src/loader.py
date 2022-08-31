
from docx.text.paragraph import Paragraph
import re

Paragraph.text = property(lambda self: GetParagraphText(self))

def GetTag(element):
    return "%s:%s" % (element.prefix, re.match("{.*}(.*)", element.tag).group(1))

def GetParagraphText(paragraph):
    text = ''
    runCount = 0
    for child in paragraph._p:
        tag = GetTag(child)
        if tag == "w:r":
            text += paragraph.runs[runCount].text
            runCount += 1
        if tag == "w:hyperlink":
            for subChild in child:
                if GetTag(subChild) == "w:r":
                    text += f'{LINK_BEG}{subChild.text}{LINK_END}'
    return text


from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# path to the calendar docx
DOC_PATH = 'docx/calendar.docx'
# number of columns in the calendar
COLS  = 5
# list of months in Spanish
MONTHS = ['enero'     ,     
          'febrero'   ,  
          'marzo'     ,    
          'abril'     ,    
          'mayo'      ,    
          'junio'     ,     
          'julio'     ,    
          'agosto'    ,   
          'septiembre', 
          'octubre'   ,   
          'noviembre' , 
          'diciembre']

# the link insertion delimiter
LINK_DEL = '@LINK'
LINK_BEG = '<a href=@LINK>'
LINK_END = '</a>'

class PropErr: 
    '''A class for storage 
    of propagated errors''' 
    
    def __init__(self):
        self.errs = []
        
    def __str__(self):
        err_str = ""
        # concatenate errors into String
        for i, err in enumerate(self.errs):
            err_str += f'Error[{i + 1}]: {err}'
            
        return err_str
        
    def push(self, err):
        self.errs.append(err)
        
    def has_err(self):
        return len(self.errs) != 0


def month_index(month):
    '''Returns the index of a month name 
    in the MONTHS list. If the month isn't found, 
    a value of -1 i returned'''
    
    for i, month_cur in enumerate(MONTHS):
        if month_cur == month:
            return i     
    return -1

def month_days(month):
    '''Returns the amount of 
    days in a month, given its name'''
    
    match month_index(month) + 1:
        case 1:  return 31
        case 2:  return 28
        case 3:  return 31 
        case 4:  return 30
        case 5:  return 31 
        case 6:  return 30 
        case 7:  return 31 
        case 8:  return 31 
        case 9:  return 30 
        case 10: return 31 
        case 11: return 30 
        case 12: return 31

def count_of(sub, src):
    '''Returns the indices of each
    Substring pattern which appears in src'''
    
    count = 0
    index = 0
    
    while True:
        # search for sub from index
        index = src.find(sub, index)
        # if not found
        if index == -1:
            return count
        
        index += 1   
        count += 1

def enumerate_doc(doc):
    '''A generator for the docx Document 
    type. Iterates through all the table 
    cells in the docx'''
    
    for table in doc.tables:   
        for row in table.rows:
            for cell in row.cells:
                yield cell                

def date_to_id(err, date_str):
    '''Converts a String in the '#-month/#-day'
    format into 'month/#-day' and asserts the
    passed date is valid in format'''
    
    date = date_str.split('/')
    
    if len(date) > 2:
        err.push(f'input "{date_str}" has too many "/"')
        return None, None
    
    try:
        month = int(date[0])
        day   = int(date[1])
    except:
        err.push(f'incorrect input format "{date_str}"')
        return None, None
    
    if month > 12 or month < 1:
        err.push(f'incorrect month input "{month}"')
        return None, None
    
    month = MONTHS[month - 1]
    m_days = month_days(month)
    
    if day > m_days or day < 1:
        err.push(f'{month} dates restricted to 1-{m_days}, given "{day}"')
        return None, None
    
    # return month/day
    return month, day

def calendar_beg(doc):
    '''Extracts the first calendar 
    date returning it in the 
    'month/#-day' format'''
    
    dates = [] 
    # extract first COLS date days
    for cell in enumerate_doc(doc):             
        day = cell_day(cell)
        # append day number
        if day.isnumeric():
            dates.append(int(day))      
            if len(dates) == COLS: 
                break
       
    month = None
    # extract the starting month
    for cell in enumerate_doc(doc):
        if cell.text in MONTHS:
            month = month_index(cell.text)
            break
        
    # whether month dates carry
    # over from the previous month
    carries = (max(dates) - min(dates) > 4)         
    if carries: month -= 1
    
    return MONTHS[month], dates[0]

def cell_day(cell, to_int=False):
    '''Returns the day String of a
    given cell, returning an empty
    String if there's no date'''
    
    day = ""
    # extract day number
    for ch in cell.text:
        if ch.isdigit(): 
            day += ch
        else: break
    
    if not to_int:
        return day
    else:
        return int(day)

def format_cell(doc, cell, link_index):
    '''Formats a given cell's text'''
    
    rels = doc.part.rels
    links = [rels[rel]._target for rel in rels 
             if rels[rel].reltype == RT.HYPERLINK]
    
    for relId, rel in doc.part.rels.items():
        if rel.reltype == RT.HYPERLINK:
            print(relId)
            print(rel._target)
        
    # transform cell text into paragraphs
    text = (par.text for par in cell.paragraphs)
    res  = ""
    # format paragraph text
    for part in text:
        # split source along link insertion
        split = part.split(LINK_DEL)
                
        concat = split[0]
        # concatenate source + links
        for i, splice in enumerate(split[1:]):
            concat += links[link_index + i] + splice
        
        res += f'\t{concat}\n'
        
    return res
    

def search_calendar(err, month, day):
    '''Find the cell of the specified date'''
    
    # load the calendar .docx
    doc = Document(DOC_PATH)   
    # the absolute link index 
    link_index = 0
    # query the first calendar date 
    month_i, day_i = calendar_beg(doc)
    # create doc generator    
    gen = enumerate_doc(doc)
    # offset generator to calendar beginning
    for cell in gen:
        link_index += count_of(LINK_BEG, cell.text)
        if cell.text in MONTHS: break
    
    for cell in gen:
        if cell.text in MONTHS: continue      
        if not cell_day(cell).isnumeric(): continue
        # return if both day and month match     
        if day_i == day and month_i == month:
            return format_cell(doc, cell, link_index)
        
        # increment link index
        link_index += count_of(LINK_BEG, cell.text)
        # collect the current day String value
        day_cur = cell_day(cell, to_int=True)
        # delta between the previous and  current day
        dif = abs(day_cur - day_i)
        # check whether to increment day or month
        if day_i == month_days(month_i) or dif > 1:
            # increment index in range of 0-11
            index = (month_index(month_i) + 1) % len(MONTHS)
            # increment month
            month_i = MONTHS[index]
            # reset day
            day_i   = 1
        else:
            # increment day
            day_i += 1
                       
    err.push(f'no information found for "{month} {day}"')
    return None                                                     

def get_date_info(date):
    err = PropErr()
    
    month, day = date_to_id(err, date)
    
    if err.has_err():
        return err
    
    text = search_calendar(err, month, day)  
    
    if err.has_err():
        return err
    
    return text
    
s = str(get_date_info('8/29'))
print(s)